import streamlit as st
import pandas as pd
from docx import Document
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

st.title("Créateur de bordereau très cool")

uploaded_files = st.file_uploader("Selectionner des fichiers pdfs", type=["pdf"], accept_multiple_files=True)

st.divider()

df_list = []


if uploaded_files:
    st.write("Files téléchargés:")
    for i, file in enumerate(uploaded_files):
        pdf_name = os.path.basename(file.name)
        pdf_name_cleaned = pdf_name.replace(".pdf", "").replace("_", " ")
        df_list.append({"file_name": pdf_name_cleaned})
    files_pd_dataframe = pd.DataFrame(df_list)
    st.dataframe(files_pd_dataframe, use_container_width=True)
    st.divider()
    button_bordereau = st.button("Créer le bordereau", use_container_width=True)

    if button_bordereau:
        document = Document()
        # Titre du document Word
        section = document.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run("BORDEREAU DE PIÈCES")
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)  # Noir
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i, row in files_pd_dataframe.iterrows():
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            font.color.rgb = RGBColor(0, 0, 0)  # Black
            words = row["file_name"].split(" ")

            # Add the words to the paragraph with the first two in bold
            for i, word in enumerate(words):
                run = paragraph.add_run(word + ' ')
                font = run.font
                if i < 2:
                    font.bold = True

            # Add a new line between each line
            paragraph.add_run().add_break()

        bureau = os.path.join(os.path.expanduser("~"), "Desktop")

        # Vérifier si le répertoire existe, sinon le créer
        if not os.path.exists(bureau):
            os.makedirs(bureau)

        # Enregistrer le fichier Word sur le bureau
        document.save(os.path.join(bureau, 'Bordereau.docx'))
        st.success('Bordereau créé : Bordereau.docx sur le bureau')

        st.success(f"Chemin complet du bureau :{os.path.join(bureau, 'Bordereau.docx')}")

        document.save('Bordereau.docx')

        st.balloons()
